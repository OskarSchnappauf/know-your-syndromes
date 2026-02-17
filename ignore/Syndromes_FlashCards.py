import docx
doc = docx.Document('/Users/oskar/Python/small_projects/Syndromes_Flash_Cards/Essential genetic syndromes.docx')
all_paras = doc.paragraphs
syndromes = []
for paragraph in doc.paragraphs:
    #print(paragraph.style)
    if paragraph.style.name=='Heading 1':
        syndromes.append(paragraph.text)
for i, element in enumerate(syndromes):
    print(i,element)